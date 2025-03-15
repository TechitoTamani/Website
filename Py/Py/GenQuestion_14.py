
def Q14(section,document,Ansdoc):
  from docx import Document
  import random

  question = "section ขณะที่บอลลูนลอยขึ้นเป็นแนวเส้นตรงในแนวดิ่งด้วยความเร็วคงที่ u m/s ที่ความสูง s เมตร คนอยู่ในบอลลูนได้ทำขวดตกลงมา จงหาว่าขวดจะใช้เวลาตกลงมานานเท่าใดจึงถึงพื้น".split()

  # QuestionCode

  u = random.choice([i*10 for i in range (1,11)])
  s = random.randint(1,100)
    
  while ((2*s)%10) != 0 or (((2*s)/10)**(1/2)) != int(((2*s)/10)**(1/2)):
    s = random.randint(1,100)

  for i in range (len(question)) :
    if question[i] == "section":
          question[i] = "".join([str(section),")"])
    elif question[i] == "u":
        question[i] = str(u)
    elif question[i] == "s":
        question[i] = str(s)

  # AnsCode

  Ans = ((2*s)/10)**(1/2)

  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20]

  Num = random.randint(0,(len(MinusPlusType)-1))

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice and WrongChoice > 0 :
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))

  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q14(i,test1)