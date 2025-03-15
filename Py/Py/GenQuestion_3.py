
def Q3(section,document,Ansdoc):

  from docx import Document
  import random

  question = "section รถคันหนึ่งวิ่งด้วยความเร็ว u1 km/hr ต่อมาความเร็วของรถเปลี่ยนเป็น u2 km/hr ความเร็วเฉลี่ยของรถคันนี้เป็นกี่ m/s".split()

  # QuestionCode

  u1 = random.randint(1,100)
  u2 = random.randint(1,100)
    
  while (((u1*5)/18)+((u2*5)/18))%2 != 0 :
    u1 = random.randint(1,100)
    u2 = random.randint(1,100)

  for i in range (len(question)) :
      if question[i] == "section":
          question[i] = "".join([str(section),")"])
      elif question[i] == "u1":
          question[i] = str(u1)
      elif question[i] == "u2":
          question[i] = str(u2)

  # AnsCode

  Ans = (((u1*5)/18)+((u2*5)/18))/2


  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20]

  Num = random.randint(0,(len(MinusPlusType)-1))

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice and WrongChoice > 0 :
      choice.append(WrongChoice)


  # PrintCode
  
  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))

  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q3(i,test1)