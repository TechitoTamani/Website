
def Q6(section,document,Ansdoc):

  from docx import Document
  import random

  question = "section รถวิ่งแนวตรงด้วยความเร็วคงตัว u m/s เมื่อเบรกในทันทีจะหยุดได้ในเวลา t วินาที จงหาระยะทางที่รถเคลื่อนที่นับตั้งแต่เบรกจนกระทั่งหยุดพอดี".split()

  # QuestionCode

  u = random.randint(1,100)
  t = random.randint(1,10)
    
  while (u*t)%2 != 0 :
    u = random.randint(1,100)
    t = random.randint(1,10)

  for i in range (len(question)) :
    if question[i] == "section":
          question[i] = "".join([str(section),")"])
    elif question[i] == "u":
        question[i] = str(u)
    elif question[i] == "t":
        question[i] = str(t)

  # AnsCode

  Ans = (u*t)/2


  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20]

  Num = random.randint(0,(len(MinusPlusType)-1))

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice and WrongChoice > 0 :
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))

  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q6(i,test1)