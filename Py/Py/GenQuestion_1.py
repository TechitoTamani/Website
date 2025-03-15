
def Q1(section,document,Ansdoc):

  from docx import Document
  import random

  namelist = ["ชิ","วิน","สายฟ้า","โฟ","ต้นตาล","ภีม"]

  question = "section นาย name วิ่งทางตรงเป็นระยะทาง s1 เมตร ใช้เวลา t1 วินาที แล้ววิ่งกลับทางเดิม s2 เมตร ใช้เวลา t2 วินาที จึงหยุด จงหาอัตราเร็วเฉลี่ยและความเร็วเฉลี่ยในหน่วยเมตรต่อวินาทีตามลำดับ".split()

  # QuestionCode

  name = namelist[random.randint(0,(len(namelist)-1))]

  s1 = random.randint(1,10)*10
  t1 = random.randint(1,10)
  s2 = random.randint(1,10)*10
  t2 = random.randint(1,10)

  while ((s1+s2)%(t1+t2)) != 0 or ((s1-s2)%(t1+t2)) != 0 :
    s1 = random.randint(1,10)*10
    t1 = random.randint(1,10)
    s2 = random.randint(1,10)*10
    t2 = random.randint(1,10)

  for i in range (len(question)) :
      if question[i] == "section":
          question[i] = "".join([str(section),")"])
      elif question[i] == "name":
          question[i] = name
      elif question[i] == "s1":
          question[i] = str(s1)
      elif question[i] == "t1":
          question[i] = str(t1)
      elif question[i] == "s2":
          question[i] = str(s2)
      elif question[i] == "t2":
          question[i] = str(t2)


  # AnsCode

  Ans1 = int((s1+s2)/(t1+t2))
  Ans2 = int((s1-s2)/(t1+t2))
  
  Ans = [Ans1,Ans2]

  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  while len(choice) < choiceAmount :

    WrongChoice1 = choice[random.randint(0,(len(choice)-1))][0]+(1*((-1)**(random.randint(0,1))))

    while WrongChoice1 < 0 :
       WrongChoice1 = choice[random.randint(0,(len(choice)-1))][0]+(1*((-1)**(random.randint(0,1))))

    WrongChoice2 = choice[random.randint(0,(len(choice)-1))][1]+(1*((-1)**(random.randint(0,1))))

    WrongChoice = [WrongChoice1,WrongChoice2]

    if WrongChoice not in choice :
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),", ".join([str(choice[i][0]), str(choice[i][1])])]))

  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))

  # document.save("test.docx")

# for i in range (1,10) :
#     Q1(i,test1)