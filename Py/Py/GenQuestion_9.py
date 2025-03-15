
def Q9(section,document,Ansdoc):

  from docx import Document
  import random

  question = "section วัตถุเคลื่อนที่จากหยุดนิ่งด้วยความเร่งคงตัว a m/s^2 เมื่อเวลาผ่านไป t วินาทีจะได้การกระจัดกี่เมตร".split()

  # QuestionCode

  a = random.randint(1,50)
  t = random.randint(1,10)
    
  while (a*(t**2))%2 != 0 :
    a = random.randint(1,50)
    t = random.randint(1,10)

  for i in range (len(question)) :
    if question[i] == "section":
          question[i] = "".join([str(section),")"])
    elif question[i] == "a":
        question[i] = str(a)
    elif question[i] == "t":
        question[i] = str(t)

  # AnsCode

  Ans = (a*(t**2))/2


  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20,100]

  Num = random.randint(0,4)

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  if Ans >= 100:
     Num = 5

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice and WrongChoice > 0:
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))

  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q9(i,test1)