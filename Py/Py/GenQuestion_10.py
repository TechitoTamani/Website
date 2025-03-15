
def Q10(section,document,Ansdoc):

  from docx import Document
  import random
  
  namelist = ["ชิ","วิน","สายฟ้า","โฟ","ต้นตาล","ภีม"]

  question = "section นาย name วิ่งออกจากจุดสตาร์ทด้วยควรามเร่งคงที่ตลอดทางเท่ากับ a m/s^2 โดยวิ่งถึงเส้นชัยใช้เวลา t วินาที นาย name เข้าเส้นชัยด้วยความเร็วเท่าใด".split()

  # QuestionCode

  name = namelist[random.randint(0,(len(namelist)-1))]

  a = random.randint(1,50)
  t = random.randint(1,10)

  while a*t > 100 :
    a = random.randint(1,50)
    t = random.randint(1,10)

  for i in range (len(question)) :
    if question[i] == "section":
          question[i] = "".join([str(section),")"])
    elif question[i] == "a":
        question[i] = str(a)
    elif question[i] == "t":
        question[i] = str(t)
    elif question[i] == "name":
        question[i] = name

  # AnsCode

  Ans = a*t

  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20]

  Num = random.randint(0,(len(MinusPlusType)-1))

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice and WrongChoice > 0 :
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))

  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q10(i,test1)