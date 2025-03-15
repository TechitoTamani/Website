
def Q11(section,document,Ansdoc):

  from docx import Document
  import random

  question = "section วัตถุกำลังเคลื่อนที่ในแนวตรงด้วยความเร่งคงที่ ในช่วง 2 วินาทีติดๆ กัน วัตถุเคลื่อนที่ได้ระยะทาง s1 เมตรกับ s2 เมตรตามลำดับ ความเร่งของการเคลื่อนที่เป็นเท่าใด".split()

  # QuestionCode

  s1 = random.randint(1,10)
  s2 = random.randint(1,10)
    
  while ((s2/2)-(s1))%2 != 0 :
    s1 = random.randint(1,10)
    s2 = random.randint(1,10)

  for i in range (len(question)) :
    if question[i] == "section":
          question[i] = "".join([str(section),")"])
    elif question[i] == "s1":
        question[i] = str(s1)
    elif question[i] == "s2":
        question[i] = str(s2)

  # AnsCode

  u = s1/1
  v = s2/2

  Ans = (v-u)/2


  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20]

  Num = random.randint(0,(len(MinusPlusType)-1))

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice:
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))

  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q11(i,test1)