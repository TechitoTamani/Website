
def Q8(section,document,Ansdoc):

  from docx import Document
  import random

  namelist = ["ชิ","วิน","สายฟ้า","โฟ","ต้นตาล","ภีม"]

  question = "section นาย name ขับรถโดยเร่งจากความเร็ว u1 km/hr เป็น u2 km/hr ในเวลา t วินาที ความเร่งเฉลี่ยของรถมีค่ากี่ m/s^2".split()

  # QuestionCode

  name = namelist[random.randint(0,(len(namelist)-1))]

  u1 = random.randint(1,100)
  u2 = random.randint(1,100)
  t = random.randint(2,10)
    
  while (((u2*5)/18)-((u1*5)/18))%t != 0 or u2 <= u1 :
    u1 = random.randint(1,100)
    u2 = random.randint(1,100)
    t = random.randint(2,10)

  for i in range (len(question)) :
    if question[i] == "section":
          question[i] = "".join([str(section),")"])
    elif question[i] == "u1":
        question[i] = str(u1)
    elif question[i] == "u2":
        question[i] = str(u2)
    elif question[i] == "t":
        question[i] = str(t)
    elif question[i] == "name":
        question[i] = str(name)

  # AnsCode

  Ans = (((u2*5)/18)-((u1*5)/18))/t


  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20]

  Num = random.randint(0,(len(MinusPlusType)-1))

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice and WrongChoice > 0:
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))

  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q8(i,test1)