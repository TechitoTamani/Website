
#อยากทำช้อยทศนิยม

def Q13(section,document,Ansdoc):
  from docx import Document
  import random

  question = "section โยนวัตถุขึ้นในแนวดิ่งจากพื้นด้วยความเร็วต้น u m/s จะต้องใช้เวลานานเท่าใดจึงจะตกกลับมาถึงที่เดิม".split()

  # QuestionCode

  u = random.randint(1,50)
    
  while (-2*u)%(-10) != 0 :
    u = random.randint(1,50)

  for i in range (len(question)) :
    if question[i] == "section":
          question[i] = "".join([str(section),")"])
    elif question[i] == "u":
        question[i] = str(u)

  # AnsCode

  Ans = (-2*u)/(-10)

  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20]

  Num = random.randint(0,(len(MinusPlusType)-1))

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice:
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))
  
  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q13(i,test1)