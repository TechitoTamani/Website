import random
from docx import Document

import GenQuestion_ex
import GenQuestion_1
import GenQuestion_2
import GenQuestion_3
import GenQuestion_4
import GenQuestion_5
import GenQuestion_6
import GenQuestion_7
import GenQuestion_8
import GenQuestion_9
import GenQuestion_10
import GenQuestion_11
import GenQuestion_12
import GenQuestion_13
import GenQuestion_14

section = int(input())
namein = str(input())

document = Document()
Ansdoc = Document()
score = 0

for i in range (1,section+1) :

    Num = random.randint(1,14)

    if Num == 1 :
        GenQuestion_1.Q1(i,document,Ansdoc)

    elif Num == 2 :
        GenQuestion_2.Q2(i,document,Ansdoc)
    
    elif Num == 3 :
        GenQuestion_3.Q3(i,document,Ansdoc)
    
    elif Num == 4 :
        GenQuestion_4.Q4(i,document,Ansdoc)

    elif Num == 5 :
        GenQuestion_5.Q5(i,document,Ansdoc)
    
    elif Num == 6 :
        GenQuestion_6.Q6(i,document,Ansdoc)
    
    elif Num == 7 :
        GenQuestion_7.Q7(i,document,Ansdoc)
    
    elif Num == 8 :
        GenQuestion_8.Q8(i,document,Ansdoc)

    elif Num == 9 :
        GenQuestion_9.Q9(i,document,Ansdoc)
    
    elif Num == 10 :
        GenQuestion_10.Q10(i,document,Ansdoc)
    
    elif Num == 11 :
        GenQuestion_11.Q11(i,document,Ansdoc)

    elif Num == 12 :
        GenQuestion_12.Q12(i,document,Ansdoc)
    
    elif Num == 13 :
        GenQuestion_13.Q13(i,document,Ansdoc)
    
    elif Num == 14 :
        GenQuestion_14.Q14(i,document,Ansdoc)

    if i % 4 == 0 and i>0 :
        para = document.add_paragraph('')
        para.paragraph_format.page_break_before = True
    

name = "".join([namein,".","docx"])
nameAns = "".join([namein,"Ans",".","docx"])
document.save(name)
Ansdoc.save(nameAns)


# print(f"{score}/{section}")