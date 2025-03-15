
def Q5(section,document,Ansdoc):

  from docx import Document
  import random

  namelist = ["ชิ","วิน","สายฟ้า","โฟ","ต้นตาล","ภีม"]

  question = "section วัตถุเคลื่อนที่แนวตรงจากสภาพหยุดนิ่งด้วยความเร่งคงตัว a m/s^2 ที่วินาทีที่ t วัตถุจะเคลื่อนที่ได้ระยะทางเท่าใด และมีความเร็วเท่าใด ตามลำดับ".split()

  # QuestionCode

  a = random.randint(1,10)
  t = random.randint(1,10)

  while (a*(t**2))%2 != 0 :
    a = random.randint(1,10)
    t = random.randint(1,10)

  for i in range (len(question)) :
      if question[i] == "section":
          question[i] = "".join([str(section),")"])
      elif question[i] == "a":
          question[i] = str(a)
      elif question[i] == "t":
          question[i] = str(t)

  # AnsCode

  Ans1 = int((a*(t**2))/2)
  Ans2 = int(a*t)
  
  Ans = [Ans1,Ans2]

  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  while len(choice) < choiceAmount :

    WrongChoice1 = choice[random.randint(0,(len(choice)-1))][0]+(1*((-1)**(random.randint(0,1))))

    while WrongChoice1 < 0 :
       WrongChoice1 = choice[random.randint(0,(len(choice)-1))][0]+(1*((-1)**(random.randint(0,1))))

    WrongChoice2 = choice[random.randint(0,(len(choice)-1))][1]+(1*((-1)**(random.randint(0,1))))

    WrongChoice = [WrongChoice1,WrongChoice2]

    if WrongChoice not in choice :
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),", ".join([str(choice[i][0]), str(choice[i][1])])]))
  
  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,5) :
#     Q5(i,test1)