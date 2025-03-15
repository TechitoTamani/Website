
def Q2(section,document,Ansdoc):

  from docx import Document
  import random

  question = "section วัตถุหนึ่งเริ่มเคลื่่อนที่จากหยุดนิ่ง หลังผ่านไป t วินาที อัตราเร็วของวัตถุเป็น v m/s วัตถุนี้เคลื่อนที่ด้วยอัตราเร็วเฉลี่ยเท่าใด".split()

  # QuestionCode

  t = random.randint(1,10)
  v = random.randint(1,100)
    
  while (v%2) != 0 :
    t = random.randint(1,10)
    v = random.randint(1,100)

  for i in range (len(question)) :
      if question[i] == "section":
          question[i] = "".join([str(section),")"])
      elif question[i] == "t":
          question[i] = str(t)
      elif question[i] == "v":
          question[i] = str(v)

  # AnsCode

  Ans = v/2


  # RandomAnsCode

  choice = [Ans]

  choiceAmount = 4

  MinusPlusType = [1,2,5,10,20]

  Num = random.randint(0,(len(MinusPlusType)-1))

  if Ans % 5 != 0:
     Num = random.randint(0,1)

  while len(choice) < choiceAmount :

    WrongChoice = choice[random.randint(0,(len(choice)-1))]+(MinusPlusType[Num]*((-1)**(random.randint(0,1))))

    if WrongChoice not in choice and WrongChoice > 0 :
      choice.append(WrongChoice)


  # PrintCode

  questionDoc = " ".join(question)

  document.add_paragraph(questionDoc)

  choice.sort()

  for i in range (choiceAmount):
    document.add_paragraph(" ".join(["".join([str(i+1),"."]),str(choice[i])]))
  
  for i in range (choiceAmount) :
     if choice[i] == Ans :
        Ansdoc.add_paragraph(" ".join(["".join([str(section),")"]),str(i+1)]))
  # document.save("test.docx")

# for i in range (1,3) :
#     Q2(i,test1)